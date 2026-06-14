from pathlib import Path

from docx import Document
from pypdf import PdfReader


def read_docx_text(file_path):
    document = Document(str(file_path))
    text_parts = []

    for paragraph in document.paragraphs:
        text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_values = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_values.append(cell_text)
            if row_values:
                text_parts.append("\n".join(row_values))

    text = "\n".join(text_parts).strip()

    if not text:
        raise ValueError("DOCX не содержит текста для импорта.")

    return text


def read_pdf_text(file_path):
    reader = PdfReader(str(file_path))
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            text_parts.append(page_text)

    text = "\n".join(text_parts).strip()

    if not text:
        raise ValueError("PDF не содержит извлекаемого текстового слоя.")

    return text


def read_document_text(file_path):
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == ".docx":
        return read_docx_text(path)

    if extension == ".pdf":
        return read_pdf_text(path)

    raise ValueError("Поддерживаются только файлы DOCX и текстовые PDF.")
