from docx import Document

from app.base_scanner import (
    clean_cell_value,
    extract_key_values_from_lines,
    extract_key_values_from_table,
    finalize_scan_result,
    make_empty_scan_result,
)


def scan_docx(file_path):
    document = Document(str(file_path))
    result = make_empty_scan_result("docx")

    paragraphs = []
    text_parts = []

    for paragraph in document.paragraphs:
        paragraph_text = clean_cell_value(paragraph.text)
        if paragraph_text:
            paragraphs.append(paragraph_text)
            text_parts.append(paragraph_text)

    result["paragraphs"] = paragraphs
    result["key_values"].update(extract_key_values_from_lines(paragraphs))

    for table_index, table in enumerate(document.tables):
        table_rows = []

        for row in table.rows:
            cells = [clean_cell_value(cell.text) for cell in row.cells]
            if any(cells):
                table_rows.append(cells)
                text_parts.append(" | ".join(cell for cell in cells if cell))

        if table_rows:
            result["tables"].append(table_rows)
            result["key_values"].update(
                extract_key_values_from_table(table_rows, table_index)
            )

    result["plain_text"] = "\n".join(text_parts).strip()

    if not result["plain_text"]:
        raise ValueError("DOCX не содержит текста для импорта.")

    return finalize_scan_result(result)
