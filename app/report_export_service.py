from __future__ import annotations

import base64
import re
import unicodedata
from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from urllib.parse import unquote, urlparse
import xml.etree.ElementTree as ET

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from flask import current_app
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as PdfImage,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from werkzeug.utils import secure_filename

from app.html_sanitizer import sanitize_editor_html
from app.models import ImportedDataBlock, ReportEditorState


EXPORT_FORMATS = {"docx", "pdf", "xml"}

CYRILLIC_TRANSLIT = {
    "а": "a",
    "б": "b",
    "в": "v",
    "г": "g",
    "д": "d",
    "е": "e",
    "ё": "e",
    "ж": "zh",
    "з": "z",
    "и": "i",
    "й": "y",
    "к": "k",
    "л": "l",
    "м": "m",
    "н": "n",
    "о": "o",
    "п": "p",
    "р": "r",
    "с": "s",
    "т": "t",
    "у": "u",
    "ф": "f",
    "х": "h",
    "ц": "c",
    "ч": "ch",
    "ш": "sh",
    "щ": "sch",
    "ъ": "",
    "ы": "y",
    "ь": "",
    "э": "e",
    "ю": "yu",
    "я": "ya",
}


@dataclass
class ReportExportData:
    report: object
    draft: object
    title: str
    report_date: date
    tag: str
    template_title: str
    author_name: str
    author_id: int | None
    source: str
    created_at: datetime | None
    updated_at: datetime | None
    document_html: str
    blocks: list


def build_report_export_data(report, current_user=None):
    draft = _find_draft_for_report(report.id)
    editor_state = _get_editor_state(draft.id) if draft else None
    document_html = sanitize_editor_html(editor_state.document_html if editor_state else "")
    blocks = _get_draft_blocks(draft.id) if draft else []
    template_title = draft.template_title if draft else ""

    return ReportExportData(
        report=report,
        draft=draft,
        title=report.report_title or "Без названия",
        report_date=report.report_date,
        tag=report.tag or "",
        template_title=template_title or report.template_key or "Не выбран",
        author_name=(report.report_author or "").strip()
        or (getattr(current_user, "name", "") or "").strip()
        or (getattr(current_user, "username", "") or "").strip()
        or "Пользователь",
        author_id=getattr(current_user, "id", None),
        source=report.source_filename or report.source_type or "manual",
        created_at=report.created_at,
        updated_at=report.updated_at,
        document_html=document_html,
        blocks=blocks,
    )


def build_export_filename(export_data, extension):
    date_part = _format_date_for_filename(export_data.report_date)
    slug = _slugify(export_data.title)
    filename = "report_{0}_{1}_{2}.{3}".format(
        export_data.report.id,
        slug or "report",
        date_part,
        extension,
    )
    return secure_filename(filename) or "report_{0}.{1}".format(export_data.report.id, extension)


def export_report_to_docx(export_data):
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    title = document.add_heading(export_data.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_docx_metadata(document, export_data)

    content_nodes = _extract_content_nodes(export_data)
    if content_nodes:
        for node_index, node in enumerate(content_nodes):
            if isinstance(node, PageMarker):
                document.add_page_break()
                continue
            _append_docx_block(document, node)
    else:
        document.add_paragraph("Содержимое отчета пока не заполнено.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def export_report_to_pdf(export_data):
    _register_pdf_fonts()
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=export_data.title,
        author=export_data.author_name,
    )
    styles = _build_pdf_styles()
    story = [
        Paragraph(_escape_pdf_text(export_data.title), styles["Title"]),
        Spacer(1, 0.35 * cm),
    ]
    story.extend(_build_pdf_metadata(export_data, styles))
    story.append(Spacer(1, 0.4 * cm))

    content_nodes = _extract_content_nodes(export_data)
    if content_nodes:
        for node in content_nodes:
            if isinstance(node, PageMarker):
                story.append(PageBreak())
                continue
            story.extend(_append_pdf_block(node, styles))
    else:
        story.append(Paragraph("Содержимое отчета пока не заполнено.", styles["Normal"]))

    doc.build(story)
    buffer.seek(0)
    return buffer


def export_report_to_xml(export_data):
    root = ET.Element("report", {"schema_version": "1.0", "id": str(export_data.report.id)})
    metadata = ET.SubElement(root, "metadata")
    ET.SubElement(metadata, "title").text = export_data.title
    ET.SubElement(metadata, "report_date").text = _format_date_iso(export_data.report_date)
    ET.SubElement(metadata, "tag").text = export_data.tag
    template = ET.SubElement(metadata, "template")
    ET.SubElement(template, "id").text = str(export_data.draft.id) if export_data.draft else ""
    ET.SubElement(template, "name").text = export_data.template_title
    author = ET.SubElement(metadata, "author")
    ET.SubElement(author, "id").text = str(export_data.author_id or "")
    ET.SubElement(author, "name").text = export_data.author_name
    ET.SubElement(metadata, "source").text = export_data.source
    ET.SubElement(metadata, "created_at").text = _format_datetime_iso(export_data.created_at)
    ET.SubElement(metadata, "updated_at").text = _format_datetime_iso(export_data.updated_at)
    ET.SubElement(root, "content", {"format": "html"}).text = export_data.document_html or _blocks_to_html(export_data.blocks)

    tree = ET.ElementTree(root)
    buffer = BytesIO()
    tree.write(buffer, encoding="utf-8", xml_declaration=True)
    buffer.seek(0)
    return buffer


def _find_draft_for_report(report_id):
    from app.models import ReportDraft

    return (
        ReportDraft.query.filter_by(linked_report_id=report_id)
        .order_by(ReportDraft.updated_at.desc(), ReportDraft.created_at.desc())
        .first()
    )


def _get_editor_state(draft_id):
    return ReportEditorState.query.filter_by(draft_id=draft_id).first()


def _get_draft_blocks(draft_id):
    blocks = (
        ImportedDataBlock.query.filter_by(draft_id=draft_id, is_deleted=False)
        .order_by(ImportedDataBlock.order_index.asc())
        .all()
    )
    return [
        {
            "type": block.block_type,
            "content_text": block.content_text or "",
            "content_json": block.content_json or "",
            "source_file_name": block.source_file_name or "",
        }
        for block in blocks
    ]


def _add_docx_metadata(document, export_data):
    rows = [
        ("Дата отчета", _format_date_display(export_data.report_date)),
        ("Тег", export_data.tag or "Не указан"),
        ("Шаблон", export_data.template_title or "Не выбран"),
        ("Автор", export_data.author_name),
        ("Источник", export_data.source or "manual"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    document.add_paragraph()


def _extract_content_nodes(export_data):
    html = export_data.document_html or _blocks_to_html(export_data.blocks)
    soup = BeautifulSoup(sanitize_editor_html(html), "html.parser")
    pages = soup.select(".editable-page, [data-editable-page]")

    if pages:
        nodes = []
        for page_index, page in enumerate(pages):
            if page_index:
                nodes.append(PageMarker())
            nodes.extend(_meaningful_children(page))
        return nodes

    return _meaningful_children(soup)


def _meaningful_children(node):
    children = []
    for child in node.children:
        if isinstance(child, NavigableString):
            if child.strip():
                children.append(child)
            continue

        if isinstance(child, Tag):
            if _is_service_node(child):
                continue
            children.append(child)

    return children


def _is_service_node(node):
    service_classes = {
        "editor-page-number",
        "table-resize-handle",
        "editor-caret-marker",
        "dashboard-preview-loader",
    }
    classes = set(node.get("class") or [])
    return bool(classes & service_classes) or node.has_attr("data-page-orientation-toggle")


class PageMarker:
    pass


def _append_docx_block(document, node, list_kind=None):
    if isinstance(node, NavigableString):
        text = str(node).strip()
        if text:
            document.add_paragraph(text)
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower()

    if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = min(int(tag[1]), 4)
        paragraph = document.add_heading(level=level)
        _append_docx_inline(paragraph, node)
        return

    if tag in {"p", "div", "section", "article", "blockquote", "figcaption"}:
        if _has_direct_block_children(node):
            for child in _meaningful_children(node):
                _append_docx_block(document, child)
            return

        paragraph = document.add_paragraph(style=_list_style(list_kind))
        _apply_docx_paragraph_style(paragraph, node)
        _append_docx_inline(paragraph, node)
        return

    if tag in {"ul", "ol"}:
        nested_list_kind = "number" if tag == "ol" else "bullet"
        for child in node.find_all("li", recursive=False):
            _append_docx_block(document, child, nested_list_kind)
        return

    if tag == "li":
        paragraph = document.add_paragraph(style=_list_style(list_kind or "bullet"))
        _append_docx_inline(paragraph, node)
        for nested_list in node.find_all(["ul", "ol"], recursive=False):
            _append_docx_block(document, nested_list)
        return

    if tag == "table":
        _append_docx_table(document, node)
        return

    if tag == "img":
        paragraph = document.add_paragraph()
        _append_docx_image(paragraph, node)
        return

    if tag == "br":
        document.add_paragraph()
        return

    paragraph = document.add_paragraph()
    _append_docx_inline(paragraph, node)


def _append_docx_inline(paragraph, node, style_context=None):
    style_context = dict(style_context or {})

    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            _apply_docx_run_style(run, style_context)
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower()
    next_context = _merge_inline_context(style_context, tag, node)

    if tag == "br":
        paragraph.add_run().add_break()
        return

    if tag == "img":
        _append_docx_image(paragraph, node)
        return

    for child in node.children:
        _append_docx_inline(paragraph, child, next_context)


def _append_docx_table(document, table_node):
    rows = table_node.find_all("tr")
    if not rows:
        return

    max_cols = max(
        sum(_safe_int(cell.get("colspan"), 1) for cell in row.find_all(["td", "th"], recursive=False))
        for row in rows
    )
    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        col_index = 0
        for cell_node in row.find_all(["td", "th"], recursive=False):
            while col_index < max_cols and table.cell(row_index, col_index).text:
                col_index += 1
            if col_index >= max_cols:
                break

            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _append_docx_inline(paragraph, cell_node)
            colspan = _safe_int(cell_node.get("colspan"), 1)

            if colspan > 1 and col_index + colspan - 1 < max_cols:
                cell.merge(table.cell(row_index, col_index + colspan - 1))

            col_index += colspan

    document.add_paragraph()


def _append_docx_image(paragraph, image_node):
    image_buffer = _get_safe_image_buffer(image_node.get("src"))
    if not image_buffer:
        return

    try:
        paragraph.add_run().add_picture(image_buffer, width=Inches(5.8))
    except Exception:
        current_app.logger.exception("Failed to add image to DOCX export")


def _apply_docx_run_style(run, context):
    run.bold = bool(context.get("bold"))
    run.italic = bool(context.get("italic"))
    run.underline = bool(context.get("underline"))

    if context.get("font_size"):
        run.font.size = Pt(context["font_size"])

    if context.get("font_family"):
        run.font.name = context["font_family"]


def _apply_docx_paragraph_style(paragraph, node):
    styles = _parse_style(node.get("style"))
    align = styles.get("text-align")
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if align in alignment_map:
        paragraph.alignment = alignment_map[align]

    if styles.get("margin-left"):
        paragraph.paragraph_format.left_indent = Cm(_css_length_to_cm(styles["margin-left"]))

    if styles.get("text-indent"):
        paragraph.paragraph_format.first_line_indent = Cm(_css_length_to_cm(styles["text-indent"]))

    if styles.get("line-height"):
        paragraph.paragraph_format.line_spacing = _css_line_height(styles["line-height"])


def _merge_inline_context(context, tag, node):
    next_context = dict(context)
    styles = _parse_style(node.get("style"))

    if tag in {"strong", "b"} or styles.get("font-weight") in {"bold", "700", "800", "900"}:
        next_context["bold"] = True

    if tag in {"em", "i"} or styles.get("font-style") == "italic":
        next_context["italic"] = True

    if tag == "u" or "underline" in styles.get("text-decoration", ""):
        next_context["underline"] = True

    font_size = _css_font_size_to_pt(styles.get("font-size"))
    if font_size:
        next_context["font_size"] = font_size

    if styles.get("font-family"):
        next_context["font_family"] = styles["font-family"].split(",")[0].strip("'\" ")

    return next_context


def _list_style(list_kind):
    if list_kind == "number":
        return "List Number"
    if list_kind == "bullet":
        return "List Bullet"
    return None


def _has_direct_block_children(node):
    block_tags = {"p", "div", "section", "article", "ul", "ol", "table", "h1", "h2", "h3", "h4", "h5", "h6"}
    return any(isinstance(child, Tag) and child.name in block_tags for child in node.children)


def _append_pdf_block(node, styles):
    if isinstance(node, NavigableString):
        text = str(node).strip()
        return [Paragraph(_escape_pdf_text(text), styles["Normal"])] if text else []

    if not isinstance(node, Tag):
        return []

    tag = node.name.lower()

    if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = min(int(tag[1]), 4)
        return [Paragraph(_node_to_pdf_inline(node), styles[f"Heading{level}"]), Spacer(1, 0.12 * cm)]

    if tag in {"p", "div", "section", "article", "blockquote", "figcaption"}:
        if _has_direct_block_children(node):
            story = []
            for child in _meaningful_children(node):
                story.extend(_append_pdf_block(child, styles))
            return story
        return [Paragraph(_node_to_pdf_inline(node), styles["Normal"]), Spacer(1, 0.1 * cm)]

    if tag in {"ul", "ol"}:
        return [_build_pdf_list(node, styles, ordered=(tag == "ol"))]

    if tag == "table":
        table = _build_pdf_table(node, styles)
        return [table, Spacer(1, 0.2 * cm)] if table else []

    if tag == "img":
        image = _build_pdf_image(node)
        return [image, Spacer(1, 0.12 * cm)] if image else []

    if tag == "br":
        return [Spacer(1, 0.1 * cm)]

    return [Paragraph(_node_to_pdf_inline(node), styles["Normal"])]


def _build_pdf_metadata(export_data, styles):
    rows = [
        ["Дата отчета", _format_date_display(export_data.report_date)],
        ["Тег", export_data.tag or "Не указан"],
        ["Шаблон", export_data.template_title or "Не выбран"],
        ["Автор", export_data.author_name],
        ["Источник", export_data.source or "manual"],
    ]
    table = Table(rows, colWidths=[4 * cm, 11 * cm], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "ExportFont"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#475569")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f8fafc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return [table]


def _build_pdf_list(list_node, styles, ordered=False):
    items = []
    for child in list_node.find_all("li", recursive=False):
        items.append(ListItem(Paragraph(_node_to_pdf_inline(child), styles["Normal"])))
    return ListFlowable(items, bulletType="1" if ordered else "bullet", leftIndent=18)


def _build_pdf_table(table_node, styles):
    rows = []
    for row in table_node.find_all("tr"):
        cells = []
        for cell in row.find_all(["td", "th"], recursive=False):
            cells.append(Paragraph(_node_to_pdf_inline(cell), styles["TableCell"]))
        if cells:
            rows.append(cells)

    if not rows:
        return None

    max_cols = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]
    table = Table(normalized_rows, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "ExportFont"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94a3b8")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f8fafc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def _build_pdf_image(image_node):
    image_buffer = _get_safe_image_buffer(image_node.get("src"))
    if not image_buffer:
        return None

    try:
        image = PdfImage(image_buffer)
        max_width = 15 * cm
        if image.drawWidth > max_width:
            ratio = max_width / image.drawWidth
            image.drawWidth = max_width
            image.drawHeight *= ratio
        return image
    except Exception:
        current_app.logger.exception("Failed to add image to PDF export")
        return None


def _node_to_pdf_inline(node, context=None):
    context = dict(context or {})

    if isinstance(node, NavigableString):
        return _escape_pdf_text(str(node))

    if not isinstance(node, Tag):
        return ""

    tag = node.name.lower()

    if tag == "br":
        return "<br/>"

    if tag == "img":
        return ""

    child_html = "".join(_node_to_pdf_inline(child, context) for child in node.children)

    if tag in {"strong", "b"}:
        return "<b>{0}</b>".format(child_html)
    if tag in {"em", "i"}:
        return "<i>{0}</i>".format(child_html)
    if tag == "u":
        return "<u>{0}</u>".format(child_html)

    styles = _parse_style(node.get("style"))
    if styles.get("font-weight") in {"bold", "700", "800", "900"}:
        child_html = "<b>{0}</b>".format(child_html)
    if styles.get("font-style") == "italic":
        child_html = "<i>{0}</i>".format(child_html)
    if "underline" in styles.get("text-decoration", ""):
        child_html = "<u>{0}</u>".format(child_html)

    return child_html


def _build_pdf_styles():
    base = getSampleStyleSheet()
    styles = {
        "Title": ParagraphStyle(
            "ExportTitle",
            parent=base["Title"],
            fontName="ExportFont-Bold",
            fontSize=20,
            leading=25,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "Normal": ParagraphStyle(
            "ExportNormal",
            parent=base["Normal"],
            fontName="ExportFont",
            fontSize=11,
            leading=15,
            alignment=TA_LEFT,
            spaceAfter=5,
        ),
        "TableCell": ParagraphStyle(
            "ExportTableCell",
            parent=base["Normal"],
            fontName="ExportFont",
            fontSize=9,
            leading=12,
        ),
    }

    for level in range(1, 5):
        styles[f"Heading{level}"] = ParagraphStyle(
            f"ExportHeading{level}",
            parent=base[f"Heading{level}"],
            fontName="ExportFont-Bold",
            fontSize=max(13, 20 - level * 2),
            leading=max(16, 24 - level * 2),
            spaceAfter=8,
        )

    return styles


def _register_pdf_fonts():
    if "ExportFont" in pdfmetrics.getRegisteredFontNames():
        return

    regular = _find_font_file(["arial.ttf", "DejaVuSans.ttf"])
    bold = _find_font_file(["arialbd.ttf", "DejaVuSans-Bold.ttf"]) or regular
    italic = _find_font_file(["ariali.ttf", "DejaVuSans-Oblique.ttf"]) or regular
    bold_italic = _find_font_file(["arialbi.ttf", "DejaVuSans-BoldOblique.ttf"]) or bold

    if not regular:
        current_app.logger.warning("Cyrillic export font was not found; PDF may render incorrectly.")
        return

    pdfmetrics.registerFont(TTFont("ExportFont", str(regular)))
    pdfmetrics.registerFont(TTFont("ExportFont-Bold", str(bold)))
    pdfmetrics.registerFont(TTFont("ExportFont-Italic", str(italic)))
    pdfmetrics.registerFont(TTFont("ExportFont-BoldItalic", str(bold_italic)))
    pdfmetrics.registerFontFamily(
        "ExportFont",
        normal="ExportFont",
        bold="ExportFont-Bold",
        italic="ExportFont-Italic",
        boldItalic="ExportFont-BoldItalic",
    )


def _find_font_file(candidates):
    roots = [
        Path("C:/Windows/Fonts"),
        Path("/usr/share/fonts/truetype/dejavu"),
        Path("/usr/share/fonts/dejavu"),
        Path("/Library/Fonts"),
    ]

    for root in roots:
        for candidate in candidates:
            path = root / candidate
            if path.exists():
                return path

    return None


def _get_safe_image_buffer(src):
    if not src:
        return None

    src = src.strip()

    if src.lower().startswith("data:image/"):
        return _image_buffer_from_data_uri(src)

    parsed = urlparse(src)
    path = unquote(parsed.path if parsed.scheme else src)

    if not path.startswith("/static/"):
        return None

    static_root = Path(current_app.static_folder).resolve()
    candidate = (static_root / path.removeprefix("/static/")).resolve()

    if static_root not in candidate.parents and candidate != static_root:
        return None

    if not candidate.exists() or candidate.suffix.lower() not in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
        return None

    return BytesIO(candidate.read_bytes())


def _image_buffer_from_data_uri(src):
    try:
        header, payload = src.split(",", 1)
    except ValueError:
        return None

    if ";base64" not in header.lower():
        return None

    mime = header.split(";", 1)[0].lower()
    if mime not in {"data:image/png", "data:image/jpeg", "data:image/jpg", "data:image/gif", "data:image/webp"}:
        return None

    try:
        return BytesIO(base64.b64decode(payload, validate=True))
    except (ValueError, TypeError):
        return None


def _blocks_to_html(blocks):
    parts = []
    for block in blocks:
        text = block.get("content_text") or ""
        if text:
            parts.append("<p>{0}</p>".format(_escape_xml_text(text)))
    return "".join(parts)


def _parse_style(style_value):
    styles = {}
    for item in (style_value or "").split(";"):
        if ":" not in item:
            continue
        key, value = item.split(":", 1)
        styles[key.strip().lower()] = value.strip().lower()
    return styles


def _css_font_size_to_pt(value):
    if not value:
        return None
    match = re.match(r"([0-9.]+)\s*(px|pt|rem|em)?", value)
    if not match:
        return None
    number = float(match.group(1))
    unit = match.group(2) or "px"
    if unit == "pt":
        return number
    if unit == "px":
        return number * 0.75
    if unit in {"rem", "em"}:
        return number * 12
    return None


def _css_length_to_cm(value):
    match = re.match(r"(-?[0-9.]+)\s*(px|pt|cm|mm|in)?", value or "")
    if not match:
        return 0
    number = float(match.group(1))
    unit = match.group(2) or "px"
    if unit == "cm":
        return number
    if unit == "mm":
        return number / 10
    if unit == "in":
        return number * 2.54
    if unit == "pt":
        return number * 0.0352778
    return number * 0.0264583


def _css_line_height(value):
    match = re.match(r"([0-9.]+)", value or "")
    if not match:
        return None
    return float(match.group(1))


def _safe_int(value, default):
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _escape_pdf_text(value):
    return (
        str(value or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _escape_xml_text(value):
    return (
        str(value or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _format_date_display(value):
    if hasattr(value, "strftime"):
        return value.strftime("%d.%m.%Y")
    return str(value or "")


def _format_date_iso(value):
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value or "")


def _format_datetime_iso(value):
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return ""


def _format_date_for_filename(value):
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return "no-date"


def _slugify(value):
    transliterated = []
    for char in (value or "").lower():
        transliterated.append(CYRILLIC_TRANSLIT.get(char, char))
    normalized = unicodedata.normalize("NFKD", "".join(transliterated))
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^a-z0-9]+", "_", ascii_value).strip("_")
    return slug[:80] or "report"
